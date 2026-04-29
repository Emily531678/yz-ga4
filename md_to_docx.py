import argparse
from pathlib import Path

from docx import Document
from docx.shared import Pt


def md_to_docx(md_path: Path, out_path: Path) -> None:
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)

    in_code = False
    code_buf: list[str] = []

    def flush_code() -> None:
        nonlocal code_buf
        if not code_buf:
            return
        p = doc.add_paragraph("\n".join(code_buf))
        for r in p.runs:
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        code_buf = []

    for raw in lines:
        line = raw.rstrip("\n")

        if line.strip().startswith("```"):
            if in_code:
                in_code = False
                flush_code()
            else:
                in_code = True
            continue

        if in_code:
            code_buf.append(line)
            continue

        if not line.strip():
            doc.add_paragraph("")
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        stripped = line.lstrip()
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue

        if any(stripped.startswith(f"{i}. ") for i in range(1, 10)):
            doc.add_paragraph(stripped, style="List Number")
            continue

        doc.add_paragraph(line.strip())

    if in_code:
        flush_code()

    doc.save(out_path)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--in", dest="in_path", default="总报告_WebAnalytics_GA4.md")
    parser.add_argument("--out", dest="out_path", default="总报告_WebAnalytics_GA4.docx")
    args = parser.parse_args()
    md_to_docx(Path(args.in_path), Path(args.out_path))
